from docx import Document
from pathlib import Path
import re
from horology import timed


@timed(unit="min")
def split_text(content):

    chunks = re.split("(<.+?>)", content)

    return chunks


@timed(unit="min")
def create_docx(chunks, page_span, path):
    document = Document()
    p = document.add_paragraph()

    for chunk in chunks:
        if chunk and chunk[0] == "<":
            sub_text = p.add_run(chunk)
            sub_text.font.subscript = True
            # sub_text.font.bold = True
            sub_text.font.name = "Jomolhari"
        else:
            normal_text = p.add_run(chunk)
            normal_text.font.name = "Jomolhari"

    p = path.parent / "docx"
    p.mkdir(parents=True, exist_ok=True)
    out_path = path.parent / "docx" / f"{source_path.stem}_{page_span[0]}-{page_span[1]}.docx"
    document.save(str(out_path))


@timed(unit="min")
def select_span(content, page_span):
    selection = ""
    trans = "".maketrans("abcdefgxyz", "0123456789")
    start = int(page_span[0].translate(trans))
    end = int(page_span[1].translate(trans))
    for line in content.splitlines():
        if re.search("\[(\d+[abcdefgxyz])\]", line):
            page = re.search("\[(.+?)\]", line).group(1)
            current = int(page.translate(trans))
        else:
            current = 0
        if current >= start and current <= end:
            selection += f"\n\n{line}"
        elif current > end:
            break
    return selection


if __name__ == "__main__":
    vol = 67
    page_span = ["178a", "187b"]
    source_path = Path(f"data/v{vol:03}/{vol}_combined.txt")
    content = source_path.read_text(encoding="utf-8")
    selection = select_span(content, page_span)
    chunks = split_text(selection)
    create_docx(chunks, page_span, source_path)
